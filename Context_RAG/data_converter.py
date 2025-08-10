"""
Data Converter for RAG Chatbot Project
Converts PDF and DOCX files into chunked documents for AstraDB storage
Using PyMuPDF + python-docx (more stable than Docling on Windows)
"""

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import os
from pathlib import Path
import warnings

# Suppress warnings
warnings.filterwarnings("ignore", category=UserWarning)

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF using PyMuPDF"""
    try:
        doc = fitz.open(pdf_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
            text += "\n\n"  # Add spacing between pages
        
        doc.close()
        return text
    except Exception as e:
        print(f"❌ Error extracting from PDF {pdf_path}: {e}")
        return ""

def extract_text_from_docx(docx_path):
    """Extract text from DOCX using python-docx"""
    try:
        doc = DocxDocument(docx_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    except Exception as e:
        print(f"❌ Error extracting from DOCX {docx_path}: {e}")
        return ""

def document_converter():
    """
    Convert PDFs and DOCX files into chunked documents for AstraDB storage
    """
    
    print("🔄 Initializing document converter...")
    
    # Initialize text splitter for chunking
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,  # Adjust based on your embedding model
        chunk_overlap=200,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    docs = []
    data_folder = "data/"
    
    # Check if data folder exists
    if not os.path.exists(data_folder):
        print(f"❌ ERROR: '{data_folder}' folder not found!")
        print("Please create a 'data' folder and add your PDF/DOCX files")
        return []
    
    print(f"📁 Scanning '{data_folder}' for files...")
    
    # List all files first
    all_files = list(Path(data_folder).glob("*"))
    pdf_docx_files = [f for f in all_files if f.suffix.lower() in ['.pdf', '.docx']]
    
    print(f"Found {len(all_files)} total files, {len(pdf_docx_files)} are PDF/DOCX")
    
    if len(pdf_docx_files) == 0:
        print("❌ No PDF or DOCX files found in data folder!")
        print("Files found:", [f.name for f in all_files])
        return []
    
    # Process all PDF and DOCX files
    for file_path in pdf_docx_files:
        print(f"📄 Processing: {file_path.name}")
        
        try:
            # Extract text based on file type
            if file_path.suffix.lower() == '.pdf':
                full_text = extract_text_from_pdf(str(file_path))
            elif file_path.suffix.lower() == '.docx':
                full_text = extract_text_from_docx(str(file_path))
            else:
                continue
            
            if not full_text.strip():
                print(f"⚠️  No text extracted from {file_path.name}")
                continue
            
            print(f"   📝 Extracted {len(full_text)} characters")
            
            # Create chunks
            chunks = text_splitter.split_text(full_text)
            print(f"   ✂️  Created {len(chunks)} chunks")
            
            # Create Document objects for each chunk
            for i, chunk in enumerate(chunks):
                metadata = {
                    "source_file": file_path.name,
                    "file_type": file_path.suffix,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_size": file_path.stat().st_size
                }
                
                doc = Document(
                    page_content=chunk,
                    metadata=metadata
                )
                docs.append(doc)
                
        except Exception as e:
            print(f"❌ Error processing {file_path.name}: {str(e)}")
            continue
    
    print(f"✅ Total chunks created: {len(docs)}")
    return docs

# Test function
if __name__ == "__main__":
    print("🚀 Testing Document Converter (PyMuPDF + python-docx)...")
    print("="*60)
    
    try:
        documents = document_converter()
        
        if len(documents) > 0:
            print(f"\n🎉 SUCCESS! Created {len(documents)} document chunks.")
            
            # Show summary by file
            file_stats = {}
            for doc in documents:
                source = doc.metadata['source_file']
                if source not in file_stats:
                    file_stats[source] = 0
                file_stats[source] += 1
            
            print(f"\n📊 Chunks per file:")
            for filename, count in file_stats.items():
                print(f"   {filename}: {count} chunks")
            
            # Show sample
            print(f"\n📋 Sample chunk from {documents[0].metadata['source_file']}:")
            print(f"Content preview: {documents[0].page_content[:300]}...")
            
        else:
            print("\n⚠️  No documents were processed. Check your data folder and files.")
            
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()